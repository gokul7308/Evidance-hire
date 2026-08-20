import pdfplumber
import docx
from backend.models.schemas import ExtractionResult, ExtractionStatus
from backend.ingestion.file_validator import validate_file

def extract_pdf(file_path: str) -> ExtractionResult:
    """Extracts text from a PDF file."""
    try:
        with pdfplumber.open(file_path) as pdf:
            pages = pdf.pages
            if not pages:
                return ExtractionResult(
                    status=ExtractionStatus.UNKNOWN,
                    file_type=".pdf",
                    error_message="UNKNOWN - no extractable text layer detected.",
                    metadata={"pages_count": 0}
                )
            
            full_text = []
            for page in pages:
                text = page.extract_text()
                if text:
                    full_text.append(text)
            
            raw_text = "\n\n".join(full_text).strip()
            
            if not raw_text:
                return ExtractionResult(
                    status=ExtractionStatus.UNKNOWN,
                    file_type=".pdf",
                    error_message="UNKNOWN - no extractable text layer detected.",
                    metadata={"pages_count": len(pages)}
                )
                
            return ExtractionResult(
                status=ExtractionStatus.SUCCESS,
                file_type=".pdf",
                raw_text=raw_text,
                metadata={"pages_count": len(pages)}
            )
    except Exception as e:
        return ExtractionResult(
            status=ExtractionStatus.ERROR,
            file_type=".pdf",
            error_message=f"PDF extraction failed: {str(e)}"
        )

def extract_docx(file_path: str) -> ExtractionResult:
    """Extracts text from a DOCX file."""
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
                
        raw_text = "\n".join(full_text).strip()
        
        # Word docs don't typically have a concept of "no text layer" like PDFs,
        # but if it's completely empty we can still return a valid result with empty text
        # depending on what is requested. The prompt says "if empty file, handle empty file".
        
        return ExtractionResult(
            status=ExtractionStatus.SUCCESS,
            file_type=".docx",
            raw_text=raw_text,
            metadata={"paragraph_count": len(doc.paragraphs)}
        )
    except Exception as e:
        return ExtractionResult(
            status=ExtractionStatus.ERROR,
            file_type=".docx",
            error_message=f"DOCX extraction failed: {str(e)}"
        )

def extract_file(file_path: str) -> ExtractionResult:
    """
    Main entry point for extracting text from a supported resume file.
    Validates the file first, then routes to the correct extractor.
    """
    is_valid, ext, error_msg = validate_file(file_path)
    
    if not is_valid:
        return ExtractionResult(
            status=ExtractionStatus.ERROR,
            file_type=ext if ext else "unknown",
            error_message=error_msg
        )
        
    if ext == '.pdf':
        return extract_pdf(file_path)
    elif ext == '.docx':
        return extract_docx(file_path)
    else:
        return ExtractionResult(
            status=ExtractionStatus.ERROR,
            file_type=ext,
            error_message=f"Unsupported file type: {ext}"
        )
